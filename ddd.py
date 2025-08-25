import streamlit as st
import requests
import json
import uuid
import pandas as pd
import numpy as np
import io
import base64
import re
from typing import Dict, Any, List, Optional
import time
from datetime import datetime
import warnings

# Suppress warnings
warnings.filterwarnings('ignore', message='Data Validation extension is not supported')
warnings.filterwarnings('ignore', message='Mean of empty slice')
warnings.filterwarnings('ignore', message='Downcasting object dtype arrays')
warnings.filterwarnings('ignore', category=FutureWarning)
warnings.filterwarnings('ignore', category=RuntimeWarning, module='numpy')

try:
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
except (ImportError, AttributeError):
    warnings.filterwarnings('ignore', message='Unverified HTTPS request')

# Import optional libraries with error handling
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("PyPDF2 not available - PDF support disabled")

try:
    import docx
    DOCX_AVAILABLE = True
    print("python-docx available - Word support enabled")
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available - Word support disabled")

# === Configuration ===
API_URL = "https://sfassist.edagenaidev.awsdns.internal.das/api/cortex/complete"
API_KEY = "78a799ea-a0f6-11ef-a0ce-15a449f7a8b0"
APP_ID = "edadip"
APLCTN_CD = "edagnai"
MODEL = "llama3.1-70b"

class FileProcessor:
    """Handle different file types and extract content"""
    
    @staticmethod
    def process_csv(file) -> Dict[str, Any]:
        """Process CSV files and return structured data"""
        try:
            # Set pandas options to avoid warnings
            pd.set_option('future.no_silent_downcasting', True)
            
            # Try different encodings if UTF-8 fails
            try:
                df = pd.read_csv(file, encoding='utf-8')
            except UnicodeDecodeError:
                file.seek(0)  # Reset file pointer
                try:
                    df = pd.read_csv(file, encoding='latin-1')
                except UnicodeDecodeError:
                    file.seek(0)  # Reset file pointer
                    df = pd.read_csv(file, encoding='cp1252')
            
            result = {
                "type": "csv",
                "filename": file.name,
                "summary": "",
                "raw_data": {}
            }
            
            # Clean column names
            df.columns = df.columns.astype(str).str.strip()
            
            # Handle missing values properly
            df_clean = df.copy()
            for col in df_clean.columns:
                if df_clean[col].dtype == 'object':
                    df_clean[col] = df_clean[col].fillna('').astype(str)
                else:
                    df_clean[col] = df_clean[col].fillna(0)
            
            sheet_info = {
                "rows": int(len(df)),
                "columns": int(len(df.columns)),
                "column_names": [str(col) for col in df.columns.tolist()],
                "sample_data": df_clean.head(3).to_dict('records'),
                "data_types": {str(k): str(v) for k, v in df.dtypes.to_dict().items()},
                "summary_stats": {}
            }
            
            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            for col in numeric_cols:
                try:
                    valid_data = df[col].dropna()
                    if len(valid_data) > 0:
                        stats = {
                            "mean": float(valid_data.mean()),
                            "median": float(valid_data.median()),
                            "min": float(valid_data.min()),
                            "max": float(valid_data.max()),
                            "count": int(valid_data.count()),
                            "null_count": int(df[col].isnull().sum())
                        }
                        # Convert all stats to strings to avoid type issues
                        sheet_info["summary_stats"][str(col)] = {str(k): str(v) for k, v in stats.items()}
                except Exception:
                    continue
            
            result["sheet_info"] = sheet_info
            result["raw_data"]["main"] = df_clean.head(100).to_dict('records')  # Limit to 100 rows
            result["summary"] = f"CSV file '{file.name}' with {len(df)} rows and {len(df.columns)} columns"
            
            return result
            
        except Exception as e:
            return {"type": "csv", "filename": file.name, "error": str(e)}
        """Process Excel files and return structured data"""
        try:
            # Set pandas options to avoid warnings
            pd.set_option('future.no_silent_downcasting', True)
            
            excel_data = pd.read_excel(file, sheet_name=None)
            
            result = {
                "type": "excel",
                "filename": file.name,
                "sheets": {},
                "summary": "",
                "raw_data": {}
            }
            
            for sheet_name, df in excel_data.items():
                # Clean column names
                df.columns = df.columns.astype(str).str.strip()
                
                # Handle missing values properly
                df_clean = df.copy()
                for col in df_clean.columns:
                    if df_clean[col].dtype == 'object':
                        df_clean[col] = df_clean[col].fillna('').astype(str)
                    else:
                        df_clean[col] = df_clean[col].fillna(0)
                
                sheet_info = {
                    "rows": int(len(df)),
                    "columns": int(len(df.columns)),
                    "column_names": [str(col) for col in df.columns.tolist()],
                    "sample_data": df_clean.head(3).to_dict('records'),
                    "data_types": {str(k): str(v) for k, v in df.dtypes.to_dict().items()},
                    "summary_stats": {}
                }
                
                # Add summary statistics for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                for col in numeric_cols:
                    try:
                        valid_data = df[col].dropna()
                        if len(valid_data) > 0:
                            stats = {
                                "mean": float(valid_data.mean()),
                                "median": float(valid_data.median()),
                                "min": float(valid_data.min()),
                                "max": float(valid_data.max()),
                                "count": int(valid_data.count()),
                                "null_count": int(df[col].isnull().sum())
                            }
                            # Convert all stats to strings to avoid type issues
                            sheet_info["summary_stats"][str(col)] = {str(k): str(v) for k, v in stats.items()}
                    except Exception:
                        continue
                
                result["sheets"][sheet_name] = sheet_info
                result["raw_data"][sheet_name] = df_clean.head(100).to_dict('records')  # Limit to 100 rows
                
            total_rows = sum(info["rows"] for info in result["sheets"].values())
            total_sheets = len(result["sheets"])
            result["summary"] = f"Excel file '{file.name}' with {total_sheets} sheet(s) and {total_rows} total rows"
            
            return result
            
        except Exception as e:
            return {"type": "excel", "filename": file.name, "error": str(e)}
    
    @staticmethod
    def process_word(file) -> Dict[str, Any]:
        """Process Word documents and extract text"""
        if not DOCX_AVAILABLE:
            return {
                "type": "word", 
                "filename": file.name,
                "error": "python-docx library not available"
            }
            
        try:
            doc = docx.Document(file)
            
            paragraphs = []
            tables_data = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            full_text = "\n".join(paragraphs)
            
            return {
                "type": "word",
                "filename": file.name,
                "paragraphs": paragraphs[:50],  # Limit paragraphs
                "tables": tables_data,
                "full_text": full_text[:5000],  # Limit text length
                "summary": f"Word document '{file.name}' with {len(paragraphs)} paragraphs, {len(tables_data)} tables, and {len(full_text.split())} words",
                "word_count": len(full_text.split())
            }
            
        except Exception as e:
            return {"type": "word", "filename": file.name, "error": str(e)}
    
    @staticmethod
    def process_pdf(file) -> Dict[str, Any]:
        """Process PDF files and extract text"""
        if not PDF_AVAILABLE:
            return {
                "type": "pdf", 
                "filename": file.name,
                "error": "PyPDF2 library not available"
            }
            
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            pages_text = []
            full_text = ""
            
            for page_num, page in enumerate(pdf_reader.pages[:10]):  # Limit to 10 pages
                try:
                    page_text = page.extract_text()
                    pages_text.append({
                        "page": page_num + 1,
                        "text": page_text.strip()[:1000]  # Limit page text
                    })
                    full_text += page_text + "\n"
                except Exception as e:
                    pages_text.append({
                        "page": page_num + 1,
                        "text": f"Error extracting text: {str(e)}"
                    })
            
            return {
                "type": "pdf",
                "filename": file.name,
                "pages": pages_text,
                "full_text": full_text[:5000],  # Limit text length
                "summary": f"PDF document '{file.name}' with {len(pdf_reader.pages)} pages and {len(full_text.split())} words",
                "page_count": len(pdf_reader.pages),
                "word_count": len(full_text.split())
            }
            
        except Exception as e:
            return {"type": "pdf", "filename": file.name, "error": str(e)}

class SFAssistAPI:
    """Enhanced API client with better conversation handling"""
    
    def __init__(self):
        self.api_url = API_URL
        self.api_key = API_KEY
        self.app_id = APP_ID
        self.aplctn_cd = APLCTN_CD
        self.model = MODEL
    
    def create_system_message(self, files_data: List[Dict[str, Any]]) -> str:
        """Create system message with file context"""
        base_msg = """You are an expert AI Data Analyst Assistant. You analyze uploaded files and provide insights, statistics, and recommendations.

CORE RESPONSIBILITIES:
• Analyze data comprehensively across all uploaded files
• Provide detailed statistical analysis, trends, and patterns
• Generate actionable insights and recommendations  
• Answer questions accurately based on the available data
• Explain findings in clear, understandable terms

IMPORTANT: DO NOT PROVIDE ANY CODE OR PROGRAMMING SYNTAX
• Never include Python, SQL, R, or any programming code
• Never show technical implementation details
• Never provide code snippets, functions, or scripts
• Focus only on business insights, analysis, and recommendations
• Use plain English explanations and business terminology

COMMUNICATION STYLE:
• Professional business language only
• Use structured formatting with bullet points and headers
• Include specific numbers, percentages, and data points
• Provide both summaries and detailed breakdowns
• Focus on business value and actionable insights
• Use charts descriptions instead of code to create charts

RESPONSE FORMAT:
• Start with key findings summary
• Provide detailed analysis with numbers
• End with actionable recommendations
• Use business terminology, not technical jargon

"""
        
        if not files_data:
            return base_msg + "\nNo files have been uploaded yet. Please ask the user to upload files to begin analysis."
        
        context = f"\n=== UPLOADED FILES ===\nYou have access to {len(files_data)} file(s):\n\n"
        
        for i, file_data in enumerate(files_data, 1):
            if "error" in file_data:
                continue
                
            try:
                # Ensure all data is string-safe
                filename = str(file_data.get('filename', 'Unknown'))
                file_type = str(file_data.get('type', 'Unknown')).upper()
                summary = str(file_data.get('summary', 'N/A'))
                
                context += f"FILE {i}: {filename}\n"
                context += f"Type: {file_type}\n"
                context += f"Summary: {summary}\n"
                
                if file_data.get('type') == 'excel':
                    sheets = file_data.get("sheets", {})
                    for sheet_name, sheet_info in sheets.items():
                        try:
                            rows = str(sheet_info.get('rows', 0))
                            columns = str(sheet_info.get('columns', 0))
                            col_names = sheet_info.get('column_names', [])
                            col_names_str = ', '.join([str(col) for col in col_names[:5]])
                            
                            context += f"  Sheet '{str(sheet_name)}': {rows} rows × {columns} columns\n"
                            context += f"  Columns: {col_names_str}...\n"
                        except Exception:
                            context += f"  Sheet '{str(sheet_name)}': Data available\n"
                
                elif file_data.get('type') == 'csv':
                    try:
                        sheet_info = file_data.get("sheet_info", {})
                        rows = str(sheet_info.get('rows', 0))
                        columns = str(sheet_info.get('columns', 0))
                        col_names = sheet_info.get('column_names', [])
                        col_names_str = ', '.join([str(col) for col in col_names[:5]])
                        
                        context += f"  CSV Data: {rows} rows × {columns} columns\n"
                        context += f"  Columns: {col_names_str}...\n"
                        
                        # Add summary stats if available
                        summary_stats = sheet_info.get('summary_stats', {})
                        if summary_stats:
                            context += f"  Numeric columns with stats: {', '.join(summary_stats.keys())}\n"
                    except Exception:
                        context += "  CSV Data: Available\n"
                            
                elif file_data.get('type') in ['word', 'pdf']:
                    try:
                        word_count = str(file_data.get('word_count', 0))
                        context += f"Content length: {word_count} words\n"
                        
                        # Include text preview - ensure it's a string
                        full_text = str(file_data.get('full_text', ''))
                        if full_text and full_text.strip():
                            preview = full_text[:800] + "..." if len(full_text) > 800 else full_text
                            context += f"Content preview:\n{preview}\n"
                    except Exception:
                        context += "Content: Text available\n"
                
                context += "\n" + "-"*40 + "\n"
            except Exception as e:
                # Skip problematic file data
                context += f"FILE {i}: Error processing file data - {str(e)}\n"
                continue
        
        return (base_msg + context)[:6000]  # Limit total length
    
    def send_message(self, user_message: str, files_context: List[Dict[str, Any]] = None, 
                    conversation_history: List[Dict[str, str]] = None) -> Optional[str]:
        """Send message with proper error handling"""
        try:
            session_id = str(uuid.uuid4())
            
            # Create system message
            sys_msg = self.create_system_message(files_context or [])
            
            # Build messages array - ensure all are strings
            messages = []
            
            # Add recent conversation history with robust error handling
            if conversation_history:
                print(f"Processing {len(conversation_history)} conversation history items")
                for i, msg in enumerate(conversation_history[-4:]):
                    try:
                        if isinstance(msg, dict) and 'content' in msg and 'role' in msg:
                            content = msg.get('content')
                            role = msg.get('role')
                            
                            # Validate content and role are not None and can be converted to string
                            if content is not None and role is not None:
                                content_str = str(content).strip()
                                role_str = str(role).strip()
                                
                                if content_str and role_str:
                                    messages.append({
                                        "role": role_str,
                                        "content": content_str
                                    })
                                    print(f"Added history message {i}: role={role_str}, content_length={len(content_str)}")
                                else:
                                    print(f"Skipped empty history message {i}")
                            else:
                                print(f"Skipped invalid history message {i}: content={content}, role={role}")
                        else:
                            print(f"Skipped malformed history message {i}: {type(msg)}")
                    except Exception as e:
                        print(f"Error processing history message {i}: {e}")
                        continue
            
            # Add current user message
            messages.append({
                "role": "user",
                "content": str(user_message).strip()
            })
            
            # Build payload with all string values to avoid type comparison errors
            payload = {
                "query": {
                    "aplctn_cd": str(self.aplctn_cd),
                    "app_id": str(self.app_id),
                    "api_key": str(self.api_key),
                    "method": str("cortex"),
                    "model": str(self.model),
                    "sys_msg": str(sys_msg),
                    "limit_convs": str("0"),  # Ensure this is a string
                    "prompt": {
                        "messages": messages
                    },
                    "app_lvl_prefix": str("enhanced_data_analyst"),
                    "user_id": str("data_analyst_enhanced"),
                    "session_id": str(session_id)
                }
            }
            
            headers = {
                "Content-Type": "application/json; charset=utf-8",
                "Accept": "application/json",
                "Authorization": f'Snowflake Token="{str(self.api_key)}"'
            }
            
            # Debug: Print payload types to identify the issue
            print("=== PAYLOAD DEBUG ===")
            for key, value in payload["query"].items():
                print(f"{key}: {type(value).__name__} = {repr(value)}")
            print("=== MESSAGES DEBUG ===")
            for i, msg in enumerate(messages):
                print(f"Message {i}: role={type(msg['role']).__name__}({repr(msg['role'])}), content={type(msg['content']).__name__}(length={len(msg['content'])})")
            print("=====================")
            
            response = requests.post(
                self.api_url, 
                headers=headers, 
                json=payload, 
                verify=False, 
                timeout=30
            )
            
            print(f"Response status: {response.status_code}")
            if response.status_code != 200:
                print(f"Response content: {response.text}")
            
            if response.status_code == 200:
                raw = response.text
                if "end_of_stream" in raw:
                    answer, _, _ = raw.partition("end_of_stream")
                    return answer.strip()
                return raw.strip()
            else:
                error_msg = f"API Error {response.status_code}: {response.text}"
                st.error(error_msg)
                return None
                
        except Exception as e:
            print(f"Exception in send_message: {str(e)}")
            print(f"Exception type: {type(e).__name__}")
            st.error(f"Request failed: {str(e)}")
            return None

def create_quick_questions():
    """Define quick questions for different file types"""
    return {
        "📊 Data Analysis": {
            "What are the key insights from my data?": "Analyze the uploaded files and provide the top 5 key insights with supporting data points.",
            "Show me summary statistics": "Provide comprehensive summary statistics for all numeric columns.",
            "What trends do you see?": "Identify and explain the main trends, patterns, and correlations in the data.",
            "Are there any outliers or anomalies?": "Find and analyze any outliers, anomalies, or unusual patterns in the data.",
            "Compare different segments": "Compare and contrast different segments, categories, or groups in the data.",
            "What recommendations do you have?": "Based on the data analysis, provide specific actionable recommendations."
        },
        "📋 Document Analysis": {
            "Summarize the main points": "Provide a comprehensive summary of the main points, findings, and conclusions from the document.",
            "Extract key data and metrics": "Extract all numerical data, statistics, and key metrics mentioned in the document.",
            "What are the recommendations?": "Identify and summarize all recommendations, action items, and next steps mentioned in the document.",
            "Find important dates and deadlines": "Extract all important dates, deadlines, milestones, and time-sensitive information from the document."
        }
    }

def render_chat_message(role: str, content: str, timestamp: str):
    """Render individual chat messages with enhanced styling"""
    if role == "user":
        st.markdown(f"""
        <div style="display: flex; justify-content: flex-end; margin: 20px 0;">
            <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                        color: white; padding: 22px 28px; border-radius: 25px 25px 8px 25px; 
                        max-width: 75%; box-shadow: 0 4px 12px rgba(102, 126, 234, 0.3);
                        font-size: 18px; line-height: 1.6; font-weight: 500;">
                {content}
                <div style="font-size: 14px; opacity: 0.85; margin-top: 10px; text-align: right; font-weight: 400;">
                    You • {timestamp}
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    else:
        # Simple content formatting
        formatted_content = content.replace('\n', '<br>')
        
        # Handle bold and italic text safely
        try:
            formatted_content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', formatted_content)
            formatted_content = re.sub(r'\*(.*?)\*', r'<em>\1</em>', formatted_content)
            formatted_content = re.sub(r'^\s*[-•]\s*(.+)$', r'&nbsp;&nbsp;• \1', formatted_content, flags=re.MULTILINE)
        except Exception:
            # If regex fails, just use the original content
            formatted_content = content.replace('\n', '<br>')
        
        st.markdown(f"""
        <div style="display: flex; justify-content: flex-start; margin: 25px 0;">
            <div style="display: flex; align-items: flex-start; max-width: 85%;">
                <div style="background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 100%); 
                            color: white; width: 50px; height: 50px; border-radius: 50%; 
                            display: flex; align-items: center; justify-content: center; 
                            margin-right: 18px; font-size: 24px; font-weight: bold;
                            box-shadow: 0 3px 10px rgba(79, 70, 229, 0.3);">
                    🤖
                </div>
                <div style="background: white; color: #1f2937; padding: 28px 32px; 
                            border-radius: 25px 25px 25px 8px; 
                            box-shadow: 0 4px 15px rgba(0,0,0,0.1); 
                            border-left: 4px solid #4f46e5;
                            font-size: 17px; line-height: 1.7; flex: 1;">
                    <div style="margin-bottom: 18px;">
                        {formatted_content}
                    </div>
                    <div style="font-size: 13px; color: #6b7280; border-top: 1px solid #f3f4f6; 
                                padding-top: 12px; font-weight: 500;">
                        🤖 AI Data Analyst • {timestamp}
                    </div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

def main():
    """Main application with stable functionality"""
    
    # Page configuration
    st.set_page_config(
        page_title="AI Data Analyst",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Set pandas options
    pd.set_option('future.no_silent_downcasting', True)
    
    # Custom CSS
    st.markdown("""
    <style>
    .main > div {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    .chat-container {
        height: 70vh;
        overflow-y: auto;
        padding: 40px 30px;
        background: linear-gradient(to bottom, #fafafa 0%, #f0f9ff 100%);
        border-radius: 20px;
        margin-bottom: 30px;
        border: 1px solid #e5e7eb;
    }
    .chat-container::-webkit-scrollbar {
        width: 8px;
    }
    .chat-container::-webkit-scrollbar-track {
        background: #f1f5f9;
        border-radius: 10px;
    }
    .chat-container::-webkit-scrollbar-thumb {
        background: #cbd5e1;
        border-radius: 10px;
    }
    .file-upload-container {
        background: white;
        padding: 25px;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin-bottom: 25px;
        border: 1px solid #e5e7eb;
    }
    .quick-questions {
        background: white;
        padding: 20px;
        border-radius: 15px;
        box-shadow: 0 3px 8px rgba(0,0,0,0.08);
        margin: 20px 0;
        border: 1px solid #e5e7eb;
    }
    .stExpander {
        border: 2px solid #e5e7eb;
        border-radius: 12px;
        margin: 10px 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .stTextInput > div > div > input {
        padding: 18px 24px;
        border-radius: 30px;
        border: 3px solid #e5e7eb;
        font-size: 17px;
        font-weight: 500;
        background: white;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
    }
    .stTextInput > div > div > input:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 4px rgba(102, 126, 234, 0.15);
        outline: none;
    }
    .stButton > button {
        border-radius: 25px;
        border: none;
        font-weight: 700;
        font-size: 15px;
        padding: 12px 24px;
        box-shadow: 0 3px 6px rgba(0,0,0,0.1);
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(0,0,0,0.15);
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Initialize session state
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "files_data" not in st.session_state:
        st.session_state.files_data = []
    if "uploaded_files_names" not in st.session_state:
        st.session_state.uploaded_files_names = []
    
    # Initialize API client
    api_client = SFAssistAPI()
    
    # Header
    st.markdown("""
    <div style="text-align: center; padding: 30px 0; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                color: white; border-radius: 20px; margin-bottom: 30px; box-shadow: 0 8px 25px rgba(102, 126, 234, 0.3);">
        <div style="font-size: 50px; margin-bottom: 15px;">🤖</div>
        <h1 style="margin: 0; font-size: 48px; font-weight: 800;">AI Data Analyst</h1>
        <p style="margin: 15px 0 0 0; font-size: 20px; opacity: 0.95;">Upload files and chat with your data using AI</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.markdown('<div class="file-upload-container">', unsafe_allow_html=True)
        st.header("📁 Upload Files")
        
        # Determine supported file types
        supported_types = ['xlsx', 'xls', 'csv']  # Always support Excel and CSV
        type_descriptions = ["Excel (.xlsx, .xls)", "CSV (.csv)"]
        
        if DOCX_AVAILABLE:
            supported_types.extend(['docx'])
            type_descriptions.append("Word (.docx)")
        
        if PDF_AVAILABLE:
            supported_types.extend(['pdf'])
            type_descriptions.append("PDF (.pdf)")
        
        help_text = f"Supported: {', '.join(type_descriptions)}"
        
        # File uploader
        uploaded_files = st.file_uploader(
            "Choose files (multiple files supported)",
            type=supported_types,
            help=help_text,
            accept_multiple_files=True
        )
        
        # Process uploaded files
        if uploaded_files:
            current_file_names = [f.name for f in uploaded_files]
            
            if current_file_names != st.session_state.uploaded_files_names:
                st.session_state.uploaded_files_names = current_file_names
                st.session_state.files_data = []
                
                with st.spinner(f"Processing {len(uploaded_files)} file(s)..."):
                    for uploaded_file in uploaded_files:
                        file_extension = uploaded_file.name.split('.')[-1].lower()
                        
                        try:
                            if file_extension in ['xlsx', 'xls']:
                                file_data = FileProcessor.process_excel(uploaded_file)
                            elif file_extension == 'csv':
                                file_data = FileProcessor.process_csv(uploaded_file)
                            elif file_extension == 'docx':
                                file_data = FileProcessor.process_word(uploaded_file)
                            elif file_extension == 'pdf':
                                file_data = FileProcessor.process_pdf(uploaded_file)
                            else:
                                continue
                                
                            if file_data and "error" not in file_data:
                                st.session_state.files_data.append(file_data)
                        except Exception as e:
                            st.error(f"Error processing {uploaded_file.name}: {str(e)}")
                
                if st.session_state.files_data:
                    st.success(f"✅ {len(st.session_state.files_data)} file(s) processed!")
        
        # Display uploaded files
        if st.session_state.files_data:
            st.markdown("### 📋 Uploaded Files")
            for file_data in st.session_state.files_data:
                if "error" not in file_data:
                    file_icon = {
                        "excel": "📊", 
                        "csv": "📈", 
                        "word": "📄", 
                        "pdf": "📕"
                    }.get(file_data.get('type'), "📄")
                    st.markdown(f"""
                    <div style="background: linear-gradient(135deg, #e8f5e8 0%, #f0f9ff 100%); 
                                border-left: 5px solid #10b981; padding: 15px; border-radius: 10px; 
                                margin: 8px 0; font-size: 13px; box-shadow: 0 2px 4px rgba(0,0,0,0.05);">
                        <div style="font-weight: 700; font-size: 15px; color: #1f2937; margin-bottom: 8px;">
                            {file_icon} {file_data.get('filename', 'Unknown')}
                        </div>
                        <div style="color: #374151; font-size: 12px;">
                            {file_data.get('summary', 'No summary')}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Quick Questions
        st.markdown('<div class="quick-questions">', unsafe_allow_html=True)
        st.markdown("### 🎯 Quick Questions")
        
        quick_questions = create_quick_questions()
        
        for category, questions in quick_questions.items():
            with st.expander(category):
                for question, detailed_prompt in questions.items():
                    if st.button(question, key=f"q_{hash(question)}", use_container_width=True):
                        if not st.session_state.files_data:
                            st.warning("Upload files first!")
                        else:
                            with st.spinner("🤔 Analyzing..."):
                                # Get conversation history with proper type safety
                                conversation_history = []
                                try:
                                    print(f"Quick question: Processing {len(st.session_state.messages)} messages from session state")
                                    for i, (role, message, timestamp) in enumerate(st.session_state.messages[-4:]):
                                        print(f"Quick question message {i}: role={type(role).__name__}({repr(role)}), message={type(message).__name__}(len={len(str(message)) if message else 0})")
                                        if message and str(message).strip():
                                            conversation_history.append({
                                                "role": str(role), 
                                                "content": str(message).strip()
                                            })
                                    print(f"Quick question: Built conversation history with {len(conversation_history)} messages")
                                except Exception as e:
                                    print(f"Error processing conversation history in quick questions: {e}")
                                    conversation_history = []
                                
                                response = api_client.send_message(
                                    detailed_prompt,
                                    st.session_state.files_data,
                                    conversation_history
                                )
                                
                                if response:
                                    timestamp = datetime.now().strftime("%H:%M")
                                    # Store with explicit string conversion
                                    st.session_state.messages.append((str("user"), str(question), str(timestamp)))
                                    st.session_state.messages.append((str("assistant"), str(response), str(timestamp)))
                                    st.rerun()
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Library status - only show if there are issues
        missing_libraries = []
        if not DOCX_AVAILABLE:
            missing_libraries.append("📄 Word (.docx)")
        if not PDF_AVAILABLE:
            missing_libraries.append("📕 PDF (.pdf)")
            
        if missing_libraries:
            st.markdown("### ⚠️ Optional Libraries")
            st.info(f"Install these for full support: {', '.join(missing_libraries)}")
            
            if not DOCX_AVAILABLE:
                st.markdown("**Install Word support:** `pip install python-docx`")
            if not PDF_AVAILABLE:
                st.markdown("**Install PDF support:** `pip install PyPDF2`")
        else:
            # Show success status
            st.markdown("### ✅ All Libraries Available")
            st.success("Excel, CSV, Word, and PDF support fully enabled!")
    
    # Main chat area
    chat_container = st.container()
    
    with chat_container:
        # Only show welcome message when NO messages exist
        if not st.session_state.messages:
            if not st.session_state.files_data:
                st.markdown("""
                <div style="text-align: center; padding: 80px 50px; background: white; 
                            border-radius: 25px; margin: 40px 0; 
                            box-shadow: 0 10px 30px rgba(0,0,0,0.1); border: 1px solid #e5e7eb;">
                    <div style="font-size: 60px; margin-bottom: 20px;">🤖</div>
                    <h1 style="color: #1f2937; font-size: 32px; font-weight: 800; margin-bottom: 20px; 
                               background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                               -webkit-background-clip: text; -webkit-text-fill-color: transparent; 
                               background-clip: text;">Welcome to AI Data Analyst!</h1>
                    <p style="font-size: 20px; color: #4b5563; margin: 25px 0; font-weight: 500;">
                        Upload your files in the sidebar to get started</p>
                    <p style="font-size: 17px; color: #6b7280; opacity: 0.9;">
                        Supports Excel, CSV, Word, and PDF files for comprehensive analysis</p>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="text-align: center; padding: 60px 50px; background: white; 
                            border-radius: 25px; margin: 40px 0; 
                            box-shadow: 0 10px 30px rgba(0,0,0,0.1); border: 1px solid #e5e7eb;">
                    <div style="font-size: 50px; margin-bottom: 15px;">🚀</div>
                    <h2 style="color: #1f2937; font-size: 28px; font-weight: 700; margin-bottom: 15px;">
                        Ready to analyze your data!</h2>
                    <p style="font-size: 18px; color: #4b5563; margin: 20px 0;">
                        Your files are uploaded and processed. Ask questions below or use quick questions in the sidebar.</p>
                    <div style="margin-top: 30px;">
                        <span style="background: #e3f2fd; padding: 10px 18px; border-radius: 25px; 
                                     margin: 0 8px; font-size: 14px; font-weight: 600; color: #1565c0;">📊 Data Analysis</span>
                        <span style="background: #f3e5f5; padding: 10px 18px; border-radius: 25px; 
                                     margin: 0 8px; font-size: 14px; font-weight: 600; color: #7b1fa2;">📋 Document Analysis</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)
        else:
            # When messages exist, show only the chat container with messages
            st.markdown('<div class="chat-container">', unsafe_allow_html=True)
            for role, message, timestamp in st.session_state.messages:
                render_chat_message(role, message, timestamp)
            st.markdown('</div>', unsafe_allow_html=True)
    
    # Chat input
    st.markdown("""
    <div style="margin: 40px 0 20px 0;">
        <h3 style="font-size: 24px; font-weight: 700; color: #1f2937; margin-bottom: 15px; 
                   display: flex; align-items: center;">
            <span style="font-size: 28px; margin-right: 10px;">💬</span> Ask about your data
        </h3>
    </div>
    """, unsafe_allow_html=True)
    
    col_input, col_send, col_clear = st.columns([5, 1, 1])
    
    with col_input:
        user_input = st.text_input(
            "Type your question here...",
            placeholder="e.g., What are the key insights from my data?",
            label_visibility="collapsed"
        )
    
    with col_send:
        send_clicked = st.button("Send 🚀", use_container_width=True)
    
    with col_clear:
        clear_clicked = st.button("Clear 🗑️", use_container_width=True)
    
    if clear_clicked:
        st.session_state.messages = []
        st.rerun()
    
    # Process input
    if send_clicked and user_input.strip():
        if not st.session_state.files_data:
            st.warning("Please upload files first to start analysis!")
        else:
            with st.spinner("🤔 Analyzing your question..."):
                # Get conversation history
                conversation_history = []
                for role, message, _ in st.session_state.messages[-4:]:
                    conversation_history.append({"role": role, "content": message})
                
                # Send to API
                response = api_client.send_message(
                    user_input, 
                    st.session_state.files_data,
                    conversation_history
                )
                
                if response:
                    timestamp = datetime.now().strftime("%H:%M")
                    st.session_state.messages.append(("user", user_input, timestamp))
                    st.session_state.messages.append(("assistant", response, timestamp))
                    st.rerun()

if __name__ == "__main__":
    main()
